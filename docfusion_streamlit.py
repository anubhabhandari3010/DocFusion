import streamlit as st
import os
import tempfile
import hashlib
import numpy as np
import docx
import openpyxl
from langchain_core.messages import HumanMessage
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_classic.chains import RetrievalQA
from langchain_groq import ChatGroq
from typing import List
from langchain_core.embeddings import Embeddings

# ============ CONFIGURATION ============
# Set your Groq API key via Streamlit secrets (for Cloud) or .streamlit/secrets.toml (for local)
# Get your key from: https://console.groq.com/keys
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except Exception:
    GROQ_API_KEY = ""
# =======================================

# Simple hash-based embeddings (no API key, no PyTorch required)
class SimpleHashEmbeddings(Embeddings):
    """
    Simple deterministic embeddings using hashing.
    This is a lightweight fallback that works without any API keys or PyTorch.
    Uses TF-IDF-like approach with character n-grams.
    """
    
    def __init__(self, dim=384):
        self.dim = dim
    
    def _text_to_vector(self, text: str) -> np.ndarray:
        """Convert text to a fixed-size vector using character n-grams and hashing."""
        vector = np.zeros(self.dim, dtype=np.float32)
        
        # Use character 3-grams and 4-grams
        for n in [3, 4]:
            for i in range(len(text) - n + 1):
                ngram = text[i:i+n].lower()
                # Hash the ngram to get an index
                hash_val = int(hashlib.md5(ngram.encode()).hexdigest(), 16)
                idx = hash_val % self.dim
                vector[idx] += 1.0
        
        # Normalize the vector
        norm = np.linalg.norm(vector)
        if norm > 0:
            vector = vector / norm
        
        return vector.tolist()
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._text_to_vector(text) for text in texts]
    
    def embed_query(self, text: str) -> List[float]:
        return self._text_to_vector(text)



def translate_text(text: str, target_language: str, model_name: str, api_key: str) -> str:
    """Translate text using Groq LLM."""
    if not text or not text.strip():
        return ""
    llm = ChatGroq(
        model_name=model_name,
        groq_api_key=api_key,
        temperature=0.3,
        max_tokens=4096
    )
    prompt = f"""Translate the following text to {target_language}.
Preserve the original formatting, paragraph structure, and meaning as closely as possible.
Do not add any explanations, notes, or commentary — output only the translated text.

Text to translate:
{text}"""
    response = llm.invoke([HumanMessage(content=prompt)])
    return response.content



def extract_documents(file_path: str, file_name: str):
    """Extract text documents from PDF, TXT, DOCX, or XLSX files."""
    ext = os.path.splitext(file_name)[1].lower()

    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
        return [Document(page_content=text, metadata={"source": file_name})]

    elif ext == ".docx":
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return [Document(page_content=text, metadata={"source": file_name})]

    elif ext == ".xlsx":
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_lines = [f"Sheet: {sheet_name}"]
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    sheet_lines.append(row_text)
            parts.append("\n".join(sheet_lines))
        text = "\n\n".join(parts)
        return [Document(page_content=text, metadata={"source": file_name})]

    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload a PDF, TXT, DOCX, or XLSX file.")


# Page configuration
st.set_page_config(page_title="DocFusion - RAG Chat", page_icon="📄", layout="wide")

# Custom CSS for better styling
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
    }
    .upload-text {
        font-size: 1.2rem;
        font-weight: 500;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'vector_db' not in st.session_state:
    st.session_state.vector_db = None
if 'qa_chain' not in st.session_state:
    st.session_state.qa_chain = None
if 'docs' not in st.session_state:
    st.session_state.docs = None
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'translated_text' not in st.session_state:
    st.session_state.translated_text = ""
if 'current_file_key' not in st.session_state:
    st.session_state.current_file_key = None
if 'processing_key' not in st.session_state:
    st.session_state.processing_key = None

# Title and description
st.title("📄 DocFusion - Document Q&A")
st.markdown("Upload a document and ask questions about its content using AI-powered RAG (Retrieval-Augmented Generation)")

# Sidebar for configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    
    # Model selection
    model_name = st.selectbox(
        "LLM Model",
        ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"],
        index=0
    )
    
    # Chunk size configuration
    chunk_size = st.slider("Chunk Size", min_value=200, max_value=1000, value=500, step=50)
    chunk_overlap = st.slider("Chunk Overlap", min_value=0, max_value=200, value=50, step=10)
    
    st.markdown("---")
    st.markdown("**How to use:**")
    st.markdown("1. Upload a document (PDF, TXT, Word, or Excel) in the right panel")
    st.markdown("2. Wait for processing (embeddings are created locally)")
    st.markdown("3. Ask questions about the document in the left chat panel")
    st.markdown("4. Translate and download the document in your chosen language")

    st.markdown("---")
    st.markdown("**Features:**")
    st.markdown("• 100% free - no paid API keys needed")
    st.markdown("• Local embeddings (no OpenAI/HuggingFace costs)")
    st.markdown("• Fast Groq LLM inference")
    st.markdown("• Supports PDF, TXT, Word, and Excel documents")
    st.markdown("• Multi-lingual AI translation with download")


# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<p class="upload-text">💬 Ask Questions</p>', unsafe_allow_html=True)
    
    # Display chat messages
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    
    # Query input
    query = st.chat_input("Ask a question about the document...")
    
    if query:
        if not st.session_state.qa_chain:
            st.error("Please set your GROQ_API_KEY in .streamlit/secrets.toml (local) or Streamlit Cloud Secrets, then upload a document first!")
        else:
            # Add user message to chat
            st.session_state.messages.append({"role": "user", "content": query})
            
            with chat_container:
                with st.chat_message("user"):
                    st.markdown(query)
                
                with st.chat_message("assistant"):
                    with st.spinner("Thinking..."):
                        try:
                            response = st.session_state.qa_chain.invoke({"query": query})["result"]
                            st.markdown(response)
                            st.session_state.messages.append({"role": "assistant", "content": response})
                        except Exception as e:
                            error_msg = f"❌ Error: {str(e)}"
                            st.error(error_msg)
                            st.session_state.messages.append({"role": "assistant", "content": error_msg})

with col2:
    st.markdown('<p class="upload-text">📤 Upload Your Document</p>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("Choose a document", type=["pdf", "txt", "docx", "xlsx"],
                                      help="Upload a PDF, TXT, Word, or Excel document")

    if uploaded_file is not None:
        file_bytes = uploaded_file.getvalue()
        file_hash = hashlib.md5(file_bytes).hexdigest()

        # Detect if a new file was uploaded
        current_file_key = (uploaded_file.name, file_hash)
        prev_file_key = st.session_state.get("current_file_key")

        if prev_file_key != current_file_key:
            st.session_state.current_file_key = current_file_key
            st.session_state.translated_text = ""
            st.session_state.messages = []

        # Determine if reprocessing is needed
        processing_key = (file_hash, chunk_size, chunk_overlap, model_name)
        needs_processing = (
            st.session_state.get("processing_key") != processing_key
            or st.session_state.get("vector_db") is None
        )

        if needs_processing:
            # Clear stale state before attempting new processing
            st.session_state.vector_db = None
            st.session_state.docs = None
            st.session_state.qa_chain = None
            st.session_state.processing_key = None

            with st.spinner("Processing document..."):
                # Save uploaded file temporarily
                file_ext = os.path.splitext(uploaded_file.name)[1]
                tmp_file_path = None
                with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                    tmp_file.write(file_bytes)
                    tmp_file_path = tmp_file.name

                try:
                    # Load and process the document
                    docs = extract_documents(tmp_file_path, uploaded_file.name)

                    if not docs:
                        st.error("❌ Could not extract any text from this document. It may be a scanned/image-based file with no embedded text.")
                        st.stop()

                    # Split into chunks
                    splitter = CharacterTextSplitter(
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap
                    )
                    chunks = splitter.split_documents(docs)

                    # Filter out empty chunks to avoid FAISS indexing errors
                    chunks = [c for c in chunks if c.page_content and c.page_content.strip()]

                    if not chunks:
                        st.error("❌ Document loaded but all content appears to be empty. The file may be scanned or image-based.")
                        st.stop()

                    # Create embeddings and vector store
                    with st.spinner("Creating embeddings and building vector store..."):
                        # Using simple hash-based embeddings (no API key, no PyTorch)
                        embeddings = SimpleHashEmbeddings(dim=384)
                        vector_db = FAISS.from_documents(chunks, embeddings)

                        st.session_state.vector_db = vector_db
                        st.session_state.docs = docs
                        st.session_state.processing_key = processing_key

                        # Create QA chain using the API key from configuration
                        if GROQ_API_KEY:
                            retriever = vector_db.as_retriever()
                            qa_chain = RetrievalQA.from_chain_type(
                                llm=ChatGroq(
                                    model_name=model_name,
                                    groq_api_key=GROQ_API_KEY
                                ),
                                chain_type="stuff",
                                retriever=retriever
                            )
                            st.session_state.qa_chain = qa_chain

                            st.success(f"✅ Document processed successfully! ({len(docs)} pages, {len(chunks)} chunks)")
                        else:
                            st.warning("⚠️ Please set your GROQ_API_KEY in .streamlit/secrets.toml (local) or Streamlit Cloud Secrets to enable Q&A")

                except Exception as e:
                    import traceback
                    st.error(f"❌ Error processing document: {type(e).__name__}: {str(e)}")
                    st.code(traceback.format_exc())
                finally:
                    # Clean up temporary file
                    if tmp_file_path and os.path.exists(tmp_file_path):
                        os.remove(tmp_file_path)

        # Lazily create QA chain if key is now available but chain is missing
        if GROQ_API_KEY and st.session_state.get("vector_db") and not st.session_state.get("qa_chain"):
            retriever = st.session_state.vector_db.as_retriever()
            st.session_state.qa_chain = RetrievalQA.from_chain_type(
                llm=ChatGroq(
                    model_name=model_name,
                    groq_api_key=GROQ_API_KEY
                ),
                chain_type="stuff",
                retriever=retriever
            )

        # ---------- Translation Section ----------
        if st.session_state.get("vector_db"):
            st.markdown("---")
            st.subheader("🌐 Translation")

            if GROQ_API_KEY:
                target_lang = st.selectbox(
                    "Target Language",
                    ["English", "Spanish", "French", "German", "Hindi", "Chinese (Simplified)",
                     "Japanese", "Portuguese", "Arabic", "Russian", "Italian", "Korean", "Dutch"],
                    key="target_lang"
                )

                if st.button("Translate Document", key="translate_btn"):
                    with st.spinner("Translating... This may take a moment for large documents."):
                        docs = st.session_state.docs
                        full_text = "\n\n".join([d.page_content for d in docs])

                        # Split into translation chunks (~3000 chars)
                        trans_chunks = []
                        current_chunk = ""
                        for paragraph in full_text.split("\n\n"):
                            if len(current_chunk) + len(paragraph) < 3000:
                                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
                            else:
                                if current_chunk:
                                    trans_chunks.append(current_chunk)
                                current_chunk = paragraph
                        if current_chunk:
                            trans_chunks.append(current_chunk)

                        translated_parts = []
                        progress_bar = st.progress(0)

                        for i, chunk in enumerate(trans_chunks):
                            translated = translate_text(chunk, target_lang, model_name, GROQ_API_KEY)
                            translated_parts.append(translated)
                            progress_bar.progress((i + 1) / len(trans_chunks))

                        st.session_state.translated_text = "\n\n".join(translated_parts)
                        progress_bar.empty()

                        st.success(f"✅ Translation to {target_lang} complete!")

                if st.session_state.get("translated_text"):
                    with st.expander("📄 View Translated Text", expanded=False):
                        st.text_area("Translation", st.session_state.translated_text, height=300)

                    file_stem = os.path.splitext(uploaded_file.name)[0]
                    st.download_button(
                        label="💾 Download Translation (.txt)",
                        data=st.session_state.translated_text.encode("utf-8"),
                        file_name=f"{file_stem}_{target_lang.lower().replace(' ', '_')}.txt",
                        mime="text/plain"
                    )
            else:
                st.info("Set your GROQ_API_KEY in secrets to enable translation.")

# Footer
st.markdown("---")
st.markdown("Built with Streamlit, LangChain, and Groq | Multi-format RAG & Translation")